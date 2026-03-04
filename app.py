import io
import json
import re
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
import streamlit as st

# Optional DOCX export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# =========================================================
# CONFIG
# =========================================================
LOOKUP_DIR = Path("./data/lookups")


# =========================================================
# HELP TEXTS (tooltip)
# =========================================================
FIELD_HELP = {
    "ID Progetto": "Identificativo univoco del progetto",
    "Titolo del Progetto": "Titolo breve e riconoscibile del progetto. E' importante che sia chiaro e conciso.",
    "Percorso alla Cartella del Progetto": "Percorso root della cartella progetto (testo libero).",
    "Descrizione del Progetto": "Descrizione sintetica (5–10 righe): obiettivo, output attesi, destinatari. Se hai difficoltà, fatti aiutare da un LLM (ChatGPT, Gemini, etc.) nella sintesi dei tuoi documenti.",
    "Stato del Progetto": "Stato del progetto (valori presi da 'stato.xlsx').",
    "Data di Inizio del Progetto": "Data di avvio formale.",
    "Data di Fine del Progetto": "Data di chiusura (o prevista).",
    
    "Committente/i": "Soggetto committente/cliente (valori presi da 'committente.xlsx').",
    "Responsabile del Progetto": "Referente interno con responsabilità del progetto (valori presi da 'team.xlsx').",
    "Gruppo di Lavoro": "Seleziona uno o più membri del team (valori presi da 'team.xlsx').",
    
    "Fonti Dati utilizzate": "Fonti dati utilizzate (una o più) per l'analisi (valori presi da 'fonti_dati.xlsx').",
    "Dettaglio Temporale dei dati": "Granularità temporale dei dati (valori presi da 'dettaglio_temporale.xlsx').",
    "Anni di Riferimento": "Anno o intervallo di riferimento dell'analisi.",
    "Strumenti di analisi utilizzati": "Strumenti e software utilizzati per l'analisi (valori presi da 'strumenti.xlsx').",
    "Livello Territoriale dei dati": "Uno o più livelli territoriali. In caso di presenza più livelli gerarchici (es. Comune, Provincia, Regione) si può scegliere di riportare quello di maggior dettaglio (valori presi da 'livello_territoriale.xlsx').",
    "Perimetro Territoriale": "Elenco sintetico dei territori coperti. Da compilare solo in caso di analisi su dati di una particolare area geografica (es. solo le province dell'Emilia-Romagna).",

    "Tipo di Output": "Tipo di output prodotti (uno o più) considerando i file di delivery del Progetto (valori presi da 'output.xlsx').",
    "Dominio del Progetto": "Ambito tematico principale (valori presi da 'dominio.xlsx').",
    "Parole Chiave": "Lista di Keywords del Progetto. Inserisci le diverse una keyword in uno dei seguenti modi: \n a) una keyword per ogni riga; \n b) separando le parole utilizzando il separatore virgola (,) o punto e virgola (;). E' importante compilare correttamente questo campo per aumentare la capacità di ricerca dell'Agente AI",
}


# =========================================================
# LOOKUPS
# =========================================================
def load_lookup_xlsx(path: Path) -> List[str]:
    df = pd.read_excel(path)
    if df.empty:
        return []
    col = df.columns[0]
    values = (
        df[col]
        .dropna()
        .astype(str)
        .map(lambda x: x.strip())
    )
    values = [v for v in values.tolist() if v]
    return sorted(list(set(values)), key=lambda s: s.lower())


@st.cache_data(show_spinner=False)
def load_all_lookups(lookup_dir: Path) -> Dict[str, List[str]]:
    lookups = {}
    if lookup_dir.exists():
        for file in sorted(lookup_dir.glob("*.xlsx")):
            lookups[file.stem] = load_lookup_xlsx(file)
    return lookups


def lk(lookups: Dict[str, List[str]], name: str, fallback: List[str]) -> List[str]:
    vals = lookups.get(name, [])
    return vals if vals else fallback


# =========================================================
# VALIDATION + EXPORT
# =========================================================
def yaml_escape(s) -> str:
    s = str(s)  # forza conversione in stringa
    s = s.replace('"', '\\"')
    return f"\"{s}\""


def to_yaml_list(values: List[str]) -> str:
    return "[" + ", ".join(yaml_escape(v) for v in values) + "]"


def build_markdown(metadata: Dict) -> str:
    order = [
        "id_progetto", "titolo", "root_path", "descrizione"
        , "stato", "data_inizio_progetto", "data_fine_progetto"

        , "committente", "responsabile_progetto", "team"

        , "fonti_dati", "dettaglio_temporale", "anno_riferimento"
        , "strumenti", "livello_territoriale", "perimetro_territoriale"
        
        , "output", "dominio", "parole_chiave"

        , "data_ultima_modifica_metadati"
    ]

    lines = ["---"]
    for k in order:
        v = metadata.get(k, "")
        if isinstance(v, list):
            lines.append(f"{k}: {to_yaml_list(v)}")
        else:
            lines.append(f"{k}: {yaml_escape(str(v))}")
    lines.append("---\n")
    return "\n".join(lines)


def build_docx(metadata: Dict) -> bytes:
    doc = Document()
    doc.add_heading("ProID – Carta di identità digitale del progetto", 1)
    for k, v in metadata.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        if isinstance(v, list):
            p.add_run(", ".join(v))
        else:
            p.add_run(str(v))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# =========================================================
# APP
# =========================================================

st.set_page_config(page_title="ProID", page_icon="img/favicon.png", layout="centered")
st.image("img/logo.png")
#st.title("ProID – Crea la Carta di Identità Digitale dei Progetti")
st.caption("""Compila i metadati e scarica la scheda in formato *Markdown* (file .md generato automaticamente da questa applicazione) e sposta successivamente questo file nella cartella di progetto.  
           Una compilazione chiara e completa della carta di identità digitale del progetto migliora significativamente la capacità dell'*assistente AI* di recuperare correttamente le informazioni nel sistema di *Knowledge Sharing*.  
           Si raccomanda quindi di dedicare qualche minuto alla compilazione dei campi.
           """)

# LOOKUPS
lookups = load_all_lookups(LOOKUP_DIR)

with st.form("proid_form", clear_on_submit=False):

    # =====================================================
    # 1) INFORMAZIONI GENERALI
    # =====================================================
    st.subheader("Informazioni generali")

    col1, col2 = st.columns([1, 2])

    with col1:
        id_progetto = st.text_input(
            "ID Progetto",
            placeholder="ECON-2025-001",
            help=FIELD_HELP["ID Progetto"],
            key="id_progetto",
        )

        titolo = st.text_input(
            "Titolo del Progetto",
            help=FIELD_HELP["Titolo del Progetto"],
            key="titolo",
        )

        root_path = st.text_input(
            "Percorso alla Cartella del Progetto",
            placeholder="N:/studi_statistici/...",
            help=FIELD_HELP["Percorso alla Cartella del Progetto"],
            key="root_path",
        )

    with col2:
        descrizione = st.text_area(
            "Descrizione del Progetto",
            height=200,
            help=FIELD_HELP["Descrizione del Progetto"],
            key="descrizione",
        )

    # RIGA 3: (1,1,1) -> stato | data_inizio | data_fine
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        options_stato = ["-- seleziona --"] + lk(lookups, "stato", [])
        stato = st.selectbox(
            "Stato del Progetto",
            options_stato,
            index=0,
            help=FIELD_HELP["Stato del Progetto"],
            key="stato",
        )
    with c2:
        data_inizio = st.date_input(
            "Data di Inizio del Progetto",
            value=None,
            help=FIELD_HELP["Data di Inizio del Progetto"],
            key="data_inizio",
        )
    with c3:
        data_fine = st.date_input(
            "Data di Fine del Progetto",
            value=None,
            help=FIELD_HELP["Data di Fine del Progetto"],
            key="data_fine",
        )

    # =====================================================
    # 2) GRUPPO DI LAVORO
    # =====================================================
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("Gruppo di lavoro")

    # RIGA 4: (1,1,1) -> committente | responsabile | team
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        committente = st.multiselect(
            "Committente/i",
            lk(lookups, "committente", ["(aggiungere valori in committente.xlsx)"]),
            help=FIELD_HELP["Committente/i"],
            key="committente",
        )
    with c2:
        responsabile_progetto = st.multiselect(
            "Responsabile del Progetto",
            lk(lookups, "team", ["(aggiungere valori in team.xlsx)"]),
            help=FIELD_HELP["Responsabile del Progetto"],
            key="responsabile_progetto",
        )
    with c3:
        team = st.multiselect(
            "Gruppo di Lavoro",
            lk(lookups, "team", []),
            help=FIELD_HELP["Gruppo di Lavoro"],
            key="team",
        )
        # Opzione sotto la riga (più leggibile)
        auto_include_resp = st.checkbox(
            "Includi il responsabile nel team",
            value=True,
            help="Se attivo, il responsabile viene aggiunto automaticameente al team se non già selezionato.",
            key="auto_include_resp",
        )

    # =====================================================
    # 3) DATI DEL PROGETTO
    # =====================================================
    st.subheader("Dati del progetto")

    # RIGA 5: (1,1,1) -> fonti | dettaglio temporale | anni riferimento
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        fonti = st.multiselect(
            "Fonti Dati utilizzate",
            lk(lookups, "fonti_dati", ["ISTAT", "Infocamere"]),
            help=FIELD_HELP["Fonti Dati utilizzate"],
            key="fonti_dati",
        )
        strumenti = st.multiselect(
            "Strumenti di analisi utilizzati",
            lk(lookups, "strumenti", ["Python", "R", "Stata"]),
            help=FIELD_HELP["Strumenti di analisi utilizzati"],
            key="strumenti",
        )
    with c2:
        dettaglio_temporale = st.multiselect(
            "Dettaglio Temporale dei dati",
            lk(lookups, "dettaglio_temporale", ["annuale", "trimestrale", "mensile"]),
            help=FIELD_HELP["Dettaglio Temporale dei dati"],
            key="dettaglio_temporale",
        )
        livello_territoriale = st.multiselect(
            "Livello Territoriale dei dati",
            lk(lookups, "livello_territoriale", ["provinciale"]),
            help=FIELD_HELP["Livello Territoriale dei dati"],
            key="livello_territoriale",
        )
    with c3:
        anno_riferimento = st.multiselect(
            "Anni di Riferimento",
            lk(lookups, "anno_riferimento", ["2024"]),
            help=FIELD_HELP["Anni di Riferimento"],
            key="anno_riferimento",
        )
        perimetro = st.text_input(
            "Perimetro Territoriale (opzionale)",
            placeholder="Province del Sud; ...",
            help=FIELD_HELP["Perimetro Territoriale"],
            key="perimetro",
        )

    # =====================================================
    # 4) CLASSIFICAZIONE DEL PROGETTO
    # =====================================================
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("Classificazione del progetto")

    # RIGA 6: (1,1,1) -> output | dominio | parole chiave
    c1, c2 = st.columns([1, 2])
    with c1:
        output = st.multiselect(
            "Tipo di Output",
            lk(lookups, "output", ["dataset", "rapporto", "slide"]),
            help=FIELD_HELP["Tipo di Output"],
            key="output",
        )
        dominio = st.multiselect(
            "Dominio del Progetto",
            lk(lookups, "dominio", ["Generale"]),
            help=FIELD_HELP["Dominio del Progetto"],
            key="dominio",
        )
    with c2:
        parole = st.text_area(
            "Parole Chiave",
            height=120,
            placeholder="Inserisci una keyword per ogni riga oppure utilizza virgola (,) o punto e virgola (;) per separarle",
            help=FIELD_HELP["Parole Chiave"],
            key="parole_chiave",
        )

    # Bottone Validazione
    col_left, col_center, col_right = st.columns([1, 1, 1])
    with col_left:
        submitted = st.form_submit_button(
            "🚀 Valida e abilita export",
            use_container_width=True
        )


if submitted:
    errors = []

    # --- helper per validazioni ---
    def is_blank_text(x) -> bool:
        return x is None or str(x).strip() == ""

    def is_empty_list(x) -> bool:
        return x is None or (isinstance(x, (list, tuple, set)) and len(x) == 0)

    def is_unselected_selectbox(x: str) -> bool:
        return x is None or str(x).strip() == "" or str(x).strip() == "-- seleziona --"

    # =========================
    # Campi obbligatori (testo)
    # =========================
    if is_blank_text(id_progetto):
        errors.append("L'ID del Progetto è obbligatorio.")

    if is_blank_text(titolo):
        errors.append("Titolo del Progetto è obbligatorio.")

    if is_blank_text(descrizione):
        errors.append("Descrizione del Progetto è obbligatoria.")

    if is_blank_text(root_path):
        errors.append("Percorso alla cartella del progetto è obbligatorio.")

    # =========================
    # Campi obbligatori (liste)
    # =========================
    if is_empty_list(committente):
        errors.append("Committente/i è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(responsabile_progetto):
        errors.append("Responsabile del Progetto è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(fonti):
        errors.append("Fonti dati è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(dettaglio_temporale):
        errors.append("Dettaglio Temporale è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(anno_riferimento):
        errors.append("Anni di Riferimento è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(livello_territoriale):
        errors.append("Livello Territoriale è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(output):
        errors.append("Tipo di Output è obbligatorio (seleziona almeno 1 valore).")

    if is_empty_list(dominio):
        errors.append("Dominio del Progetto è obbligatorio (seleziona almeno 1 valore).")

    # =========================
    # Stato (selectbox con placeholder)
    # =========================
    if is_unselected_selectbox(stato):
        errors.append("Stato del Progetto è obbligatorio (seleziona un valore).")

    # =========================
    # Team (con auto-include responsabile)
    # =========================
    team_list = list(team) if team else []
    responsabile_list = list(responsabile_progetto) if responsabile_progetto else []

    if auto_include_resp:
        for resp in responsabile_list:
            if resp not in team_list:
                team_list.append(resp)

    if len(team_list) == 0:
        errors.append("Gruppo di Lavoro è obbligatorio (seleziona almeno 1 membro).")

    # =========================
    # Date
    # =========================
    if data_inizio is None:
        errors.append("Data di Inizio del Progetto è obbligatoria.")
    if data_fine is None:
        errors.append("Data di Fine del Progetto è obbligatoria.")
    if (data_inizio is not None) and (data_fine is not None) and (data_fine < data_inizio):
        errors.append("La data di fine non può essere precedente alla data di inizio del progetto.")

    # =========================
    # Parole chiave (1 per riga)
    # =========================
    parole_list = [
        x.strip()
        for x in re.split(r"[,\n;]+", str(parole))
        if x.strip()
    ]
    if len(parole_list) == 0:
        errors.append("Parole Chiave è obbligatorio (inserisci almeno 1 keyword).")

    if errors:
        st.error("Correggi i seguenti punti:\n- " + "\n- ".join(errors))
    else:
        # Normalizzazioni coerenti con i tuoi widget (multiselect = liste)
        committente_list = list(committente) if committente else []
        responsabile_list = list(responsabile_progetto) if responsabile_progetto else []
        anno_rif_list = list(anno_riferimento) if anno_riferimento else []
        dettaglio_temp_list = list(dettaglio_temporale) if dettaglio_temporale else []
        dominio_list = list(dominio) if dominio else []

        metadata = {
            "id_progetto": id_progetto.strip(),
            "titolo": titolo.strip(),
            "root_path": root_path.strip(),                  # stringa
            "descrizione": descrizione.strip(),
            "stato": str(stato).strip(),                     # stringa (selectbox)
            "data_inizio_progetto": str(data_inizio),        # stringa ISO da date_input
            "data_fine_progetto": str(data_fine),

            "committente": committente_list,                 # lista (multiselect)
            "responsabile_progetto": responsabile_list,      # lista (multiselect)
            "team": list(team_list),                         # lista

            "fonti_dati": list(fonti),                       # lista
            "dettaglio_temporale": dettaglio_temp_list,      # lista (multiselect)
            "anno_riferimento": anno_rif_list,               # lista (multiselect)
            "strumenti": list(strumenti) if strumenti else [],# lista
            "livello_territoriale": list(livello_territoriale),  # lista
            "perimetro_territoriale": perimetro.strip(),         # stringa (opzionale)

            "output": list(output),                          # lista
            "dominio": dominio_list,                         # lista (multiselect)
            "parole_chiave": parole_list,                    # lista (una per riga)
           
            "data_ultima_modifica_metadati": date.today().isoformat(),        
            
        }

        md = build_markdown(metadata)

        col_left, col_center, col_right = st.columns([1, 1, 1])
        with col_left:
            st.success("Validazione OK. Puoi esportare i file.")
            st.download_button(
                "Scarica file in formato Markdown (.md)",
                md.encode("utf-8"),
                file_name=f"ProID_{metadata['id_progetto']}.md",
                mime="text/markdown",
            )
        
# =========================================================
# Pannello lookups (download + istruzioni)
# Metti questo blocco dopo: lookups = load_all_lookups(LOOKUP_DIR)
# =========================================================

with st.expander("📥 *Lookups*: scarica e consulta le liste dei valori disponibili", expanded=False):

    st.info(
        "Mentre compili il *form* per la creazione della carta di identità digitale, se l'opzione che ti serve non è presente nelle liste, seleziona **'ZZZ-Altro'** (se disponibile). "
        "Dopo aver generato il file **.md**, puoi aprirlo con un software di editing di testo (es. 'Blocco Note') e sostituire manualmente il valore con quello desiderato.\n\n"
        "Poi, per favore, segnala il nuovo valore a **francesco.salatesantone@tagliacarne.it**: "
        "in questo modo verrà aggiunto ai valori esistenti e sarà disponibile automaticamente anche per i tuoi colleghi."
    )

    if not LOOKUP_DIR.exists():
        st.warning(f"Cartella lookups non trovata: {LOOKUP_DIR.resolve()}")
    else:
        files = sorted(LOOKUP_DIR.glob("*.xlsx"))
        if not files:
            st.warning("Nessun file .xlsx trovato nella cartella lookups.")
        else:
            st.caption("Scarica i file Excel dei lookups (uno per volta):")

            for fp in files:
                try:
                    data = fp.read_bytes()
                except Exception as e:
                    st.error(f"Impossibile leggere {fp.name}: {e}")
                    continue

                # bottone per ciascun file
                st.download_button(
                    label=f"⬇️ {fp.stem}.xlsx",
                    data=data,
                    file_name=fp.name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"dl_lookup_{fp.stem}",
                )
